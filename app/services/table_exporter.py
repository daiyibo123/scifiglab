"""Table exporter — paper table generation (Markdown / CSV / Word)."""

import csv
import io
from typing import List, Dict, Any, Optional, Tuple
from sqlalchemy.orm import Session

from app.database.models import (
    Experiment, ExperimentGroup, ExperimentGroupItem, Metric,
)
from app.services.summary_service import get_metric_direction, summarize_experiment


def export_csv(rows: List[Dict[str, Any]]) -> str:
    """Export a list of dicts to CSV string."""
    if not rows:
        return ""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=rows[0].keys())
    writer.writeheader()
    writer.writerows(rows)
    return output.getvalue()


# ─── Paper table helpers ─────────────────────────────────────────────────

def format_metric_value(value: Optional[float], decimals: int = 4) -> str:
    if value is None:
        return "-"
    return f"{value:.{decimals}f}"


def find_best_values(
    rows: List[dict],
    metric_names: List[str],
    directions: Dict[str, str],
) -> Dict[str, Tuple[float, int]]:
    """For each metric, find the best value and its row index."""
    best: Dict[str, Tuple[float, int]] = {}
    for mname in metric_names:
        direction = directions.get(mname, "higher_better")
        for idx, row in enumerate(rows):
            val = row.get(mname)
            if val is None:
                continue
            if mname not in best:
                best[mname] = (val, idx)
            else:
                cur_best, _ = best[mname]
                if direction == "lower_better" and val < cur_best:
                    best[mname] = (val, idx)
                elif direction == "higher_better" and val > cur_best:
                    best[mname] = (val, idx)
    return best


def generate_paper_table(
    db: Session,
    group_id: int,
    user_id: int,
    metric_names: List[str],
    decimals: int = 4,
    bold_best: bool = True,
    mark_second: bool = False,
) -> Dict[str, Any]:
    """Generate paper table data for a group.

    Returns dict with keys: rows, markdown, csv_text, metric_names, directions.
    """
    group = (
        db.query(ExperimentGroup)
        .filter(ExperimentGroup.id == group_id, ExperimentGroup.user_id == user_id)
        .first()
    )
    if not group:
        return {"error": "实验组不存在"}

    items = sorted(group.items, key=lambda i: i.sort_order)
    if not items:
        return {"error": "实验组内无实验"}

    # Get directions
    directions = {mname: get_metric_direction(mname) for mname in metric_names}

    # Build rows: each row = one experiment
    rows = []
    for item in items:
        exp = db.query(Experiment).filter(Experiment.id == item.experiment_id).first()
        if not exp:
            continue
        display_name = item.display_name or exp.name
        row: Dict[str, Any] = {"_display_name": display_name, "_experiment_id": exp.id}

        # Get best value for each metric
        summary = summarize_experiment(db, exp.id, user_id)
        summary_map = {s["metric_name"]: s for s in summary}

        for mname in metric_names:
            s = summary_map.get(mname)
            row[mname] = s["best_value"] if s else None

        rows.append(row)

    if not rows:
        return {"error": "无有效实验数据"}

    # Find best values
    best_map = find_best_values(rows, metric_names, directions)

    # Find second best
    second_map: Dict[str, Tuple[float, int]] = {}
    if mark_second:
        for mname in metric_names:
            direction = directions.get(mname, "higher_better")
            best_idx = best_map.get(mname, (None, -1))[1]
            second_val = None
            second_idx = -1
            for idx, row in enumerate(rows):
                if idx == best_idx:
                    continue
                val = row.get(mname)
                if val is None:
                    continue
                if second_val is None:
                    second_val = val
                    second_idx = idx
                elif direction == "lower_better" and val < second_val:
                    second_val = val
                    second_idx = idx
                elif direction == "higher_better" and val > second_val:
                    second_val = val
                    second_idx = idx
            if second_val is not None:
                second_map[mname] = (second_val, second_idx)

    # Generate Markdown
    md_lines = []
    header = "| Method |"
    sep = "| --- |"
    for mname in metric_names:
        d = directions[mname]
        arrow = "↓" if d == "lower_better" else "↑"
        header += f" {mname} {arrow} |"
        sep += " --- |"
    md_lines.append(header)
    md_lines.append(sep)

    for idx, row in enumerate(rows):
        line = f"| {row['_display_name']} |"
        for mname in metric_names:
            val = row.get(mname)
            formatted = format_metric_value(val, decimals)
            is_best = mname in best_map and best_map[mname][1] == idx
            is_second = mname in second_map and second_map[mname][1] == idx
            if bold_best and is_best:
                formatted = f"**{formatted}**"
            elif mark_second and is_second:
                formatted = f"_{formatted}_"
            line += f" {formatted} |"
        md_lines.append(line)

    markdown = "\n".join(md_lines)

    # Generate CSV
    csv_buf = io.StringIO()
    writer = csv.writer(csv_buf)
    writer.writerow(["Method"] + metric_names)
    for row in rows:
        csv_row = [row["_display_name"]]
        for mname in metric_names:
            val = row.get(mname)
            csv_row.append(format_metric_value(val, decimals) if val is not None else "")
        writer.writerow(csv_row)
    csv_text = csv_buf.getvalue()

    return {
        "rows": [{
            "display_name": r["_display_name"],
            "experiment_id": r["_experiment_id"],
            "values": {mname: r.get(mname) for mname in metric_names},
        } for r in rows],
        "markdown": markdown,
        "csv_text": csv_text,
        "metric_names": metric_names,
        "directions": directions,
        "best_map": {k: {"value": v[0], "row_idx": v[1]} for k, v in best_map.items()},
    }


def export_word_bytes(
    rows: List[dict],
    metric_names: List[str],
    directions: Dict[str, str],
    best_map: Dict[str, dict],
    decimals: int = 4,
    bold_best: bool = True,
) -> bytes:
    """Generate a .docx file with the paper table and return raw bytes."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # Title
    title = doc.add_heading("Paper Results Table", level=2)
    title.alignment = 1  # center

    # Build header row
    header = ["Method"]
    for mname in metric_names:
        arrow = "↓" if directions.get(mname) == "lower_better" else "↑"
        header.append(f"{mname} {arrow}")

    n_cols = len(header)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header cells
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Data rows
    for i, row in enumerate(rows):
        table.rows[i + 1].cells[0].text = row["_display_name"]
        table.rows[i + 1].cells[0].paragraphs[0].runs[0].font.size = Pt(9) if table.rows[i + 1].cells[0].paragraphs[0].runs else None
        for j, mname in enumerate(metric_names):
            val = row.get(mname)
            formatted = format_metric_value(val, decimals)
            cell = table.rows[i + 1].cells[j + 1]
            cell.text = ""
            run = cell.paragraphs[0].add_run(formatted)
            run.font.size = Pt(9)
            is_best = mname in best_map and best_map[mname].get("row_idx") == i
            if bold_best and is_best:
                run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
